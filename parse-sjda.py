from docx import Document
import json
import sys
import unicodedata

doc = Document("C:/Users/Matteo/OneDrive/Documents/SJDA avec coupes.docx")


class myRun:
    def __init__(self, style, text, i, j):
        self.style = style
        self.text = text
        self.i = i
        self.j = j

    def __repr__(self):
        return f"{self.text} ({self.style})"


def getRun(i, j):
    style = doc.paragraphs[i].runs[j].style.name
    text = unicodedata.normalize(
        "NFKD", doc.paragraphs[i].runs[j].text).lstrip(".")
    if (style == "Default Paragraph Font"):
        style = doc.paragraphs[i].style.name
    if (text == "" or text == " " or text == "." or
            style == "CoUPE"):
        return False
    return myRun(style, text, i, j)


def extractRunsAndIndexes(startingParagraph, length):
    i = startingParagraph
    j = 0
    indexes = []
    extractedRuns = []
    for k in range(0, length):
        # Print progression
        print(
            f"\rExtracting run {k} on {length} : ({i},{j}) ({round(100*k/length)}%)", end="")
        indexes.append([i, j])
        try:
            newRun = getRun(i, j)
        except:
            print(f"Error on k={k}, i={i}, j={j}")
            newRun = myRun("Parsing Error",
                           "Error on paragraph " + str(i), i, j)
        finally:
            extractedRuns.append(newRun)
        if (j+1 < len(doc.paragraphs[i].runs)):
            j += 1
        elif (i+1 < len(doc.paragraphs)):
            i += 1
            j = 0
        else:
            print("Reached EoF.")
            break
    print(
        f"\rExtracting run {k+1} on {length} : ({i},{j}) ({round(100*(k+1)/length)}%)")
    return extractedRuns, indexes


def appendLine(lineToAppend):
    # Remove empty "Text" fields
    if not lineToAppend["Text"]:
        del lineToAppend["Text"]
    if not subscene:
        play[scene].append(lineToAppend)
    else:
        play[subscene].append(lineToAppend)


def printPlay():
    print("{")
    for scene in play:
        print(f" {scene} {{")
        lineIndex = 0
        for line in play[scene]:
            print(f"    {lineIndex} [")
            lineIndex += 1
            if "Error" in line:
                print("        " + line["Error"])
            if "Character" in line:
                print("        " + line["Character"])
            if "Direction" in line:
                print("        " + line["Direction"])
            if "Text" in line:
                for run in line["Text"]:
                    if len(run) >= 100:
                        run = run[:41] + " [...] " + run[-40:]
                    if run[0] == "$":
                        run = "(" + \
                            line["Inline Directions"][int(run[1])] + ")"
                    print(f"        '{run}'")


play = dict()
scene = "0"
subscene = ""
lineIndex = 0
line = {}

startingParagraph = 50
totalNumberOfRuns = 4500

allRuns, indexes = extractRunsAndIndexes(startingParagraph, totalNumberOfRuns)

totalNumberOfRuns = len(indexes)

if startingParagraph != 50:
    play[scene] = []

for k in range(totalNumberOfRuns):
    # Print progression
    print(
        f"\rProcessing run {k} on {totalNumberOfRuns} ({round(100*k/totalNumberOfRuns)}%)", end="")

    # Get new run
    [i, j] = indexes[k]
    currentRun = allRuns[k]

    # Skip empty runs
    if (not currentRun):
        continue

    # If the current run is the beginning of a new line, append the previous line
    if (currentRun.style == "Perso.Prose" or
        currentRun.style == "Perso.Vers" or
        currentRun.style == "Did.Paragraphe" or
        (k > 0 and allRuns[k-1] and allRuns[k-1].style == "Did.Paragraphe") or
        currentRun.style == "Heading 1" or
        currentRun.style == "Heading 2" or
            k == totalNumberOfRuns-1):
        if line:
            appendLine(line)
        line = {}

    # New scene
    if (currentRun.style == "Heading 1"):
        scene = currentRun.text
        subscene = ""
        play[scene] = []
        continue
    # New subscene
    if (currentRun.style == "Heading 2"):
        subscene = scene + currentRun.text.rstrip(".")
        play[subscene] = []
        continue

    # Update line

    # If current run is "Perso.xxx", add "Character" field
    if (currentRun.style == "Perso.Prose" or
            currentRun.style == "Perso.Vers"):
        line["Character"] = currentRun.text.rstrip(".").lower().capitalize()

    # If current run is "Did.Paragraphe", add "Direction" field
    if (currentRun.style == "Did.Paragraphe"):
        line["Direction"] = currentRun.text

    # If current run is "Did.Réplique", ...
    if (currentRun.style == "Did.Réplique"):
        # Let k-h be the first previous non-empty and non-"COUPE"
        h = 1
        while (k - h >= 0 and
               not allRuns[k-h] or
               allRuns[k-h].style == "COUPE"):
            h += 1
        # If this run is "Perso.xxx", add a "Direction" field ...
        previousStyle = allRuns[k-h].style
        if (previousStyle == "Perso.Prose" or
                previousStyle == "Perso.Vers"):
            line["Direction"] = currentRun.text.lstrip(" ,").rstrip(" .")
        # ... but if it is in the text, add an "Inline Directions" array
        #  and a replace the text with a placeholder
        elif (previousStyle == "Texte.Prose" or
              previousStyle == "Texte.Vers"):
            if not "Inline Directions" in line:
                line["Inline Directions"] = []
            placeholder = "$" + str(len(line["Inline Directions"]))
            line["Inline Directions"].append(currentRun.text)
            newStyle = "Texte.Prose" if previousStyle == "Texte.Prose" else "Texte.Vers"
            currentRun = myRun(newStyle, placeholder,
                               currentRun.i, currentRun.j)
            # if currentRun.j == 0:
            #     line["Text"].append(placeholder)
            # else:
            #     line["Text"][-1] += placeholder

    # Add "Text" array
    if not "Text" in line:
        line["Text"] = []

    if (currentRun.style == "Texte.Vers" or
        currentRun.style == "Texte.Prose" or
        # currentRun.style == "COUPE" or
            currentRun.style == "CORRECTION"):
        if j > 0 and line["Text"]:
            line["Text"][-1] += currentRun.text
            continue
        line["Text"].append(currentRun.text.lstrip())
    if (currentRun.style == "Parsing Error"):
        line["Error"] = ("Parsing error around here !")

print(
    f"\rProcessing run {k+1} on {totalNumberOfRuns} ({round(100*(k+1)/totalNumberOfRuns)}%)")


# Clean up empty scenes
copyOfPlay = play.copy()
for scene in copyOfPlay:
    if (not copyOfPlay[scene] or
        (len(copyOfPlay[scene]) == 1 and not copyOfPlay[scene][0])):
        del play[scene]
del copyOfPlay

# Print JSON
with open("SJDA.json", "w", encoding="utf8") as outfile:
    json.dump(play, outfile, ensure_ascii=False, indent=4)


# printPlay()
